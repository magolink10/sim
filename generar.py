import os
import random
from docx import Document

# Lista de nombres de carpetas relacionadas a un call center
nombres_carpetas = [
    "Capacitacion_Agentes",
    "Scripts_Llamadas",
    "Procedimientos_Operativos",
    "Reportes_Calidad",
    "Guiones_Ventas",
    "Manejo_Objecciones",
    "Politicas_Servicio",
    "Informes_Supervision",
    "Registros_Llamadas",
    "Documentacion_Tecnica"
]

# Posibles nombres de documentos relacionados a call center
nombres_documentos = [
    "Script_Atencion_Cliente",
    "Guia_Manejo_Objecciones",
    "Instructivo_Escalamiento",
    "Protocolo_Verificacion_Identidad",
    "Checklist_Calidad_Llamadas",
    "Guion_Ventas_Producto",
    "Informe_Supervision_Diaria",
    "Procedimiento_Soporte_Tecnico",
    "Manual_Bienvenida_Agentes",
    "Registro_Interaccion_Cliente",
    "Reporte_Tiempos_Respuesta",
    "Flujo_Solucion_Incidencias",
    "Guia_Comunicacion_Efectiva",
    "Manual_Cierre_Llamadas",
    "Protocolo_Retencion_Clientes"
]

# Contenido aleatorio típico de documentos de call center
frases_contenido = [
    "Procedimiento para manejo de llamadas entrantes.",
    "Script básico para atención al cliente.",
    "Instrucciones para escalamiento de casos.",
    "Guía de resolución de problemas comunes.",
    "Protocolo para verificación de identidad del cliente.",
    "Checklist de calidad para monitoreo de llamadas.",
    "Políticas internas para comunicación con usuarios.",
    "Flujo de trabajo para soporte técnico.",
    "Guion de bienvenida y presentación del agente.",
    "Registro de interacción con el cliente.",
    "Pasos para el manejo de clientes molestos.",
    "Indicadores clave de desempeño para agentes."
]

# Crear carpetas y archivos
for carpeta in nombres_carpetas:
    os.makedirs(carpeta, exist_ok=True)

    cantidad_archivos = random.randint(20, 30)

    for i in range(cantidad_archivos):
        # Elegir un nombre de documento aleatorio y agregar número incremental para evitar repeticiones
        nombre_base = random.choice(nombres_documentos)
        nombre_archivo = f"{nombre_base}_{i+1}.docx"

        ruta_archivo = os.path.join(carpeta, nombre_archivo)

        # Crear documento
        doc = Document()
        doc.add_heading(nombre_base.replace("_", " "), level=1)

        # Agregar entre 3 y 7 párrafos aleatorios
        for _ in range(random.randint(3, 7)):
            doc.add_paragraph(random.choice(frases_contenido))

        doc.save(ruta_archivo)

print("¡Proceso completado! Carpetas y archivos .docx generados correctamente.")
